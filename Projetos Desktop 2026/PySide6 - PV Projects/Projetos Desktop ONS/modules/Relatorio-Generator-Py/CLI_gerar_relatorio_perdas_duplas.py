import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from docx2pdf import convert
from tqdm import tqdm
import time

def adicionar_tabela_teste(doc: Document, num_linhas: int = 5, num_colunas: int = 5):
    """
    Adiciona uma tabela de teste ao documento Word para validar formatação.
    
    Args:
        doc: Documento Word (Document object)
        num_linhas: Número de linhas de dados (default 5)
        num_colunas: Número de colunas (default 5)
    """
    print(f"\n📊 Adicionando tabela de teste ({num_linhas}x{num_colunas})...")
    
    doc.add_heading('Tabela de Teste - Formatação', level=2)
    doc.add_paragraph("Esta é uma tabela de teste para validar a formatação dos relatórios:")
    
    # Cria tabela com cabeçalho + dados
    table = doc.add_table(rows=num_linhas + 1, cols=num_colunas)
    table.style = 'Table Grid'
    
    # Preenche cabeçalho
    header_cells = table.rows[0].cells
    colunas_nomes = [f"Coluna {i+1}" for i in range(num_colunas)]
    
    for j, col_name in enumerate(colunas_nomes):
        cell = header_cells[j]
        cell.text = col_name
        # Formata o cabeçalho em negrito e com fundo cinza
        paragraph = cell.paragraphs[0]
        paragraph.style = 'Heading 3'
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        # Adiciona cor de fundo (cinza escuro)
        from docx.oxml import parse_xml
        cell._element.get_or_add_tcPr().append(
            parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
        )
    
    # Preenche dados
    for i in range(num_linhas):
        row_cells = table.rows[i + 1].cells
        for j in range(num_colunas):
            row_cells[j].text = f"Dado {i+1}-{j+1}"
    
    print(f"✓ Tabela de teste adicionada com sucesso!")


def adicionar_pagina_e_tabela(doc_path: str, excel_path: str, output_docx_path: str, adicionar_tabela_teste_flag: bool = True):
    print(f"Abrindo documento Word template: {doc_path}")
    doc = Document(doc_path)
    
    # Adiciona tabela de teste se solicitado
    if adicionar_tabela_teste_flag:
        adicionar_tabela_teste(doc, num_linhas=5, num_colunas=5)
    
    # Adiciona a nova seção e configura margens
    nova_secao = doc.add_section(WD_SECTION.NEW_PAGE)
    nova_secao.top_margin = Inches(1)
    nova_secao.bottom_margin = Inches(1)
    nova_secao.left_margin = Inches(1)
    nova_secao.right_margin = Inches(1)

    doc.add_heading('Relatório de Perdas Duplas (Dados do Excel)', level=1)
    doc.add_paragraph("Abaixo estão os dados extraídos da planilha Excel:")

    # Carrega os dados do Excel
    df = pd.read_excel(excel_path)
    rows, cols = df.shape
    table = doc.add_table(rows + 1, cols)
    
    #! Mudar o layout e formatação padronizada da tabela para o ONS
    table.style = 'Table Grid'


    # Preenche o cabeçalho da tabela
    for j, column_name in enumerate(df.columns):
        table.cell(0, j).text = str(column_name)

    # Preenche os dados com barra de progresso (tqdm)
    print(f"Preenchendo tabela com {rows} linhas...")
    for i in tqdm(range(rows), desc="Progresso da Tabela"):
        for j in range(cols):
            table.cell(i + 1, j).text = str(df.values[i, j])

    doc.save(output_docx_path)
    print(f"Documento Word modificado salvo em '{output_docx_path}'.")
    return output_docx_path

def converter_docx_para_pdf(docx_path: str, pdf_path: str):
    print(f"Iniciando conversão de DOCX para PDF. Isso pode demorar MUITO...")
    start_time = time.time()
    
    try:
        # A conversão é uma operação bloqueadora. Não há progresso iterativo aqui.
        convert(docx_path, pdf_path)
        end_time = time.time()
        print(f"Conversão concluída em {end_time - start_time:.2f} segundos.")
        print(f"Arquivo PDF salvo em '{pdf_path}'.")
    except Exception as e:
        print(f"ERRO na conversão para PDF. Certifique-se que o Microsoft Office está instalado.")
        print(f"Detalhes do erro: {e}")
        raise

def gerar_relatorio_teste_tabela(template_path: str, output_docx_path: str = "relatorio_teste_tabela.docx"):
    """
    Gera um relatório de teste contendo apenas a tabela de formatação.
    Útil para validar estilos antes de adicionar dados do Excel.
    
    Args:
        template_path: Caminho do template Word
        output_docx_path: Caminho de saída do documento gerado
    """
    print(f"\n🚀 Gerando relatório de teste com tabela...")
    
    if not os.path.exists(template_path):
        print(f"❌ ERRO: Template não encontrado em '{template_path}'")
        return None
    
    try:
        doc = Document(template_path)
        
        # Adiciona tabela de teste
        adicionar_tabela_teste(doc, num_linhas=5, num_colunas=5)
        
        # Salva o documento
        doc.save(output_docx_path)
        print(f"✓ Relatório de teste gerado com sucesso em '{output_docx_path}'")
        return output_docx_path
        
    except Exception as e:
        print(f"❌ Erro ao gerar relatório de teste: {e}")
        return None


if __name__ == "__main__":
    template_word_path = r"C:\Users\pedrovictor.veras\OneDrive - Operador Nacional do Sistema Eletrico\Documentos\ESTAGIO_ONS_PVRV_2025\GitHub\Palkia-PDF-extractor\src\BulbassaurQT6-ETL\Perdas-Duplas-ETL-Desktop-V4\app\assets\docs\Lista de Contingências Duplas - Copia.docx" 
    excel_data_path = r"C:\Users\pedrovictor.veras\OneDrive - Operador Nacional do Sistema Eletrico\Documentos\ESTAGIO_ONS_PVRV_2025\GitHub\Palkia-PDF-extractor\src\BulbassaurQT6-ETL\Perdas-Duplas-ETL-Desktop-V4\app\assets\planilhas_PLC\perdas_duplas_ETL_corrigido.xlsx"
    
    output_word_path = "documento_final_gerado.docx"
    output_pdf_path = "relatorio_PLC_perdas_duplas.pdf"

    print("\n" + "="*60)
    print("🏭 GERADOR AUTOMÁTICO DE RELATÓRIOS - PERDAS DUPLAS")
    print("="*60)
    print("\nOpções:")
    print("1️⃣  - Gerar relatório COMPLETO (tabela teste + dados Excel + PDF)")
    print("2️⃣  - Gerar apenas TESTE DE TABELA (validar formatação)")
    print("3️⃣  - Sair")
    print("-"*60)
    
    opcao = input("\nEscolha uma opção (1, 2 ou 3): ").strip()
    
    if opcao == "1":
        print("\n📋 Você escolheu: Gerar Relatório Completo")
        if os.path.exists(template_word_path) and os.path.exists(excel_data_path):
            try:
                # Etapa 1: Gerar DOCX (com tabela teste + dados Excel)
                docx_gerado_path = adicionar_pagina_e_tabela(
                    template_word_path, 
                    excel_data_path, 
                    output_word_path,
                    adicionar_tabela_teste_flag=True
                )
                
                # Etapa 2: Converter para PDF
                print("\n📄 Convertendo para PDF...")
                converter_docx_para_pdf(docx_gerado_path, output_pdf_path)
                print(f"\n✅ RELATÓRIO COMPLETO GERADO COM SUCESSO!")
                print(f"   📌 Word: {output_word_path}")
                print(f"   📌 PDF:  {output_pdf_path}")
                
            except Exception as e:
                print(f"\n❌ O processamento falhou: {e}")
        else:
            print("\n❌ ERRO: Um ou ambos os arquivos de entrada não foram encontrados.")
            if not os.path.exists(template_word_path):
                print(f"   - Template não encontrado: {template_word_path}")
            if not os.path.exists(excel_data_path):
                print(f"   - Excel não encontrado: {excel_data_path}")
    
    elif opcao == "2":
        print("\n🧪 Você escolheu: Gerar Teste de Tabela")
        teste_output = "relatorio_teste_tabela.docx"
        resultado = gerar_relatorio_teste_tabela(template_word_path, teste_output)
        if resultado:
            print(f"\n✅ TESTE DE TABELA GERADO COM SUCESSO!")
            print(f"   📌 Arquivo: {resultado}")
    
    elif opcao == "3":
        print("\n👋 Saindo...")
    
    else:
        print("\n❌ Opção inválida! Por favor escolha 1, 2 ou 3.")

