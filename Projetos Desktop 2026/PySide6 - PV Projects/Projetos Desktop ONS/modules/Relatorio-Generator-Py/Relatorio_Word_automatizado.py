import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
import os
import sys
from docx2pdf import convert
from tqdm import tqdm
import time

# ==============================================================================
# ⚙️ CONFIGURAÇÕES (Edite aqui para não precisar de argumentos)
# ==============================================================================

# Caminhos absolutos ou relativos
TEMPLATE_PATH = r"C:\Users\pedrovictor.veras\OneDrive - Operador Nacional do Sistema Eletrico\Documentos\ESTAGIO_ONS_PVRV_2025\GitHub\Palkia-PDF-extractor\src\BulbassaurQT6-ETL\Perdas-Duplas-ETL-Desktop-V4\app\assets\docs\Lista de Contingências Duplas - Copia.docx"
EXCEL_PATH = r"C:\Users\pedrovictor.veras\OneDrive - Operador Nacional do Sistema Eletrico\Documentos\ESTAGIO_ONS_PVRV_2025\GitHub\Palkia-PDF-extractor\src\BulbassaurQT6-ETL\Perdas-Duplas-ETL-Desktop-V4\app\assets\planilhas_PLC\perdas_duplas_ETL_corrigido.xlsx"



SHEET_NAME = "Sheet1" # Nome da aba do Excel que será lida
OUTPUT_DOCX = "relatorio_final_gerado.docx"
OUTPUT_PDF = "relatorio_final_gerado.pdf"

# ==============================================================================

def adicionar_tabela_teste(doc: Document, num_linhas: int = 5, num_colunas: int = 5):
    """Adiciona uma tabela de teste para validar formatação."""
    print(f"📊 Adicionando tabela de teste ({num_linhas}x{num_colunas})...")
    
    doc.add_heading('Tabela de Teste - Formatação', level=2)
    doc.add_paragraph("Esta é uma tabela de teste para validar a formatação:")
    
    table = doc.add_table(rows=num_linhas + 1, cols=num_colunas)
    table.style = 'Table Grid'
    
    # Cabeçalho
    header_cells = table.rows[0].cells
    for j in range(num_colunas):
        cell = header_cells[j]
        cell.text = f"Coluna {j+1}"
        # Estilização básica
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.bold = True
        
        # Cor de fundo (XML hack para python-docx)
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading_elm = parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Dados fictícios
    for i in range(num_linhas):
        row_cells = table.rows[i + 1].cells
        for j in range(num_colunas):
            row_cells[j].text = f"Dado {i+1}-{j+1}"

def gerar_relatorio_completo():
    print(f"\n🚀 INICIANDO GERAÇÃO DE RELATÓRIO ONS")
    print(f"📂 Template: {TEMPLATE_PATH}")
    print(f"📊 Excel: {EXCEL_PATH}")

    # 1. Validação de Arquivos
    if not os.path.exists(TEMPLATE_PATH):
        print(f"❌ ERRO CRÍTICO: Template não encontrado!\n   -> {TEMPLATE_PATH}")
        sys.exit(1)
    
    if not os.path.exists(EXCEL_PATH):
        print(f"❌ ERRO CRÍTICO: Excel não encontrado!\n   -> {EXCEL_PATH}")
        sys.exit(1)

    try:
        # 2. Manipulação do Word
        doc = Document(TEMPLATE_PATH)
        
        # Adiciona tabela de teste (Opcional)
        adicionar_tabela_teste(doc)
        
        # Nova Seção para os Dados Reais
        nova_secao = doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_heading('Dados do Excel', level=1)
        
        # 3. Leitura do Excel
        print(f"📖 Lendo aba '{SHEET_NAME}' do Excel...")
        try:
            df = pd.read_excel(EXCEL_PATH, sheet_name=SHEET_NAME)
        except ValueError:
            # Tenta ler a primeira aba se o nome estiver errado
            print(f"⚠️ Aba '{SHEET_NAME}' não encontrada. Lendo a primeira aba disponível.")
            df = pd.read_excel(EXCEL_PATH)
            
        rows, cols = df.shape
        print(f"✅ Dados carregados: {rows} linhas, {cols} colunas.")

        # 4. Criação da Tabela de Dados
        table = doc.add_table(rows=rows + 1, cols=cols)
        table.style = 'Table Grid'
        
        # Cabeçalho
        for j, col_name in enumerate(df.columns):
            table.cell(0, j).text = str(col_name)
            
        # Dados (com barra de progresso)
        for i in tqdm(range(rows), desc="Preenchendo Word"):
            for j in range(cols):
                val = df.values[i, j]
                table.cell(i + 1, j).text = str(val) if pd.notna(val) else ""
        
        # Salvar DOCX
        output_abs_docx = os.path.abspath(OUTPUT_DOCX)
        doc.save(output_abs_docx)
        print(f"💾 Word salvo em: {output_abs_docx}")
        
        # 5. Conversão PDF
        output_abs_pdf = os.path.abspath(OUTPUT_PDF)
        print("🔄 Convertendo para PDF (Aguarde)...")
        convert(output_abs_docx, output_abs_pdf)
        print(f"✅ PDF gerado com sucesso: {output_abs_pdf}")

    except Exception as e:
        print(f"❌ Ocorreu um erro durante o processo: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    gerar_relatorio_completo()